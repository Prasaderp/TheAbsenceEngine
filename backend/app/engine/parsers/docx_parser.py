import io
import zipfile
import zlib
from lxml import etree
import docx
from app.engine.parsers.base import BaseParser, ParsedDocument, Section
from app.shared.errors import ValidationError

_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _parse_docx_object(doc: docx.Document) -> tuple[list[Section], str]:
    sections: list[Section] = []
    buf: list[str] = []
    current_heading = "Document"
    level = 1
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if style.startswith("Heading") and text:
            if buf:
                sections.append(Section(heading=current_heading, text="\n".join(buf), level=level))
                buf = []
            try:
                level = int(style.split()[-1])
            except ValueError:
                level = 1
            current_heading = text
        elif text:
            buf.append(text)
    if buf:
        sections.append(Section(heading=current_heading, text="\n".join(buf), level=level))
    full_text = "\n\n".join(
        f"{s.heading}\n{s.text}" if s.heading != "Document" else s.text
        for s in sections
    )
    return sections, full_text


def _fallback_xml_extract(data: bytes) -> tuple[list[Section], str]:
    """Read word/document.xml directly from the ZIP, bypassing corrupt media."""
    with zipfile.ZipFile(io.BytesIO(data), "r") as zf:
        with zf.open("word/document.xml") as xml_fh:
            tree = etree.parse(xml_fh)
    paras = tree.findall(f".//{{{_WORD_NS}}}p")
    lines: list[str] = []
    for p in paras:
        text = "".join(
            t.text or "" for t in p.findall(f".//{{{_WORD_NS}}}t")
        ).strip()
        if text:
            lines.append(text)
    full_text = "\n\n".join(lines)
    sections = [Section(heading="Document", text=full_text, level=1)] if full_text else []
    return sections, full_text


class DocxParser(BaseParser):
    async def parse(self, data: bytes, mime_type: str) -> ParsedDocument:
        try:
            doc = docx.Document(io.BytesIO(data))
            sections, full_text = _parse_docx_object(doc)
            paragraph_count = len(doc.paragraphs)
        except (zipfile.BadZipFile, zlib.error):
            try:
                sections, full_text = _fallback_xml_extract(data)
                paragraph_count = len(sections[0].text.splitlines()) if sections else 0
            except Exception as e:
                raise ValidationError(f"DOCX parsing failed: {e}") from e
        except ValidationError:
            raise
        except Exception as e:
            raise ValidationError(f"DOCX parsing failed: {e}") from e

        if not full_text.strip():
            raise ValidationError("DOCX file contains no extractable text.")
        return ParsedDocument(
            text=full_text,
            sections=sections,
            metadata={"paragraph_count": paragraph_count},
        )

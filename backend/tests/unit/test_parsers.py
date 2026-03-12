import io
import csv
import pytest
from app.engine.parsers.text_parser import TextParser
from app.engine.parsers.pdf_parser import PDFParser
from app.engine.parsers.docx_parser import DocxParser
from app.engine.parsers.csv_parser import CSVParser, XLSXParser
from app.shared.errors import ValidationError

import pdfplumber
import docx
import openpyxl


@pytest.mark.asyncio
class TestTextParser:
    async def test_plain_text(self):
        data = b"Hello world\nSecond line"
        result = await TextParser().parse(data, "text/plain")
        assert "Hello world" in result.text
        assert result.sections

    async def test_markdown_sections(self):
        data = b"# Title\nContent here\n# Another\nMore"
        result = await TextParser().parse(data, "text/markdown")
        assert len(result.sections) == 2
        assert result.sections[0].heading == "Title"
        assert result.sections[1].heading == "Another"

    async def test_empty_file(self):
        result = await TextParser().parse(b"", "text/plain")
        assert result.metadata.get("empty") is True
        assert result.text == ""

    async def test_invalid_utf8_tolerant(self):
        data = b"Valid start\xff\xfe then more"
        result = await TextParser().parse(data, "text/plain")
        assert result.text  # no exception, replacement chars used


@pytest.mark.asyncio
class TestPDFParser:
    def _make_pdf(self, text: str) -> bytes:
        import reportlab.pdfgen.canvas as canvas_mod
        buf = io.BytesIO()
        c = canvas_mod.Canvas(buf)
        c.drawString(100, 750, text)
        c.save()
        return buf.getvalue()

    async def test_empty_bytes_raises(self):
        with pytest.raises(ValidationError):
            await PDFParser().parse(b"not a pdf", "application/pdf")

    async def test_valid_pdf(self):
        try:
            data = self._make_pdf("Test document content")
            result = await PDFParser().parse(data, "application/pdf")
            assert result.text
            assert result.metadata["page_count"] >= 1
        except ImportError:
            pytest.skip("reportlab not installed")


@pytest.mark.asyncio
class TestCSVParser:
    async def test_valid_csv(self):
        buf = io.StringIO()
        w = csv.writer(buf)
        w.writerow(["name", "age", "city"])
        w.writerow(["Alice", "30", "NYC"])
        w.writerow(["Bob", "25", "LA"])
        data = buf.getvalue().encode()
        result = await CSVParser().parse(data, "text/csv")
        assert "name" in result.text
        assert "Alice" in result.text
        assert result.metadata["row_count"] == 2

    async def test_header_only_raises(self):
        data = b"col1,col2\n"
        with pytest.raises(ValidationError):
            await CSVParser().parse(data, "text/csv")

    async def test_empty_raises(self):
        with pytest.raises(ValidationError):
            await CSVParser().parse(b"", "text/csv")


@pytest.mark.asyncio
class TestXLSXParser:
    def _make_xlsx(self) -> bytes:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws.append(["product", "price", "qty"])
        ws.append(["Widget", 9.99, 100])
        ws.append(["Gadget", 24.99, 50])
        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    async def test_valid_xlsx(self):
        data = self._make_xlsx()
        result = await XLSXParser().parse(data, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        assert "product" in result.text
        assert "Widget" in result.text
        assert result.metadata["sheet_count"] == 1

    async def test_invalid_bytes_raises(self):
        with pytest.raises(ValidationError):
            await XLSXParser().parse(b"not xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@pytest.mark.asyncio
class TestDocxParser:
    def _make_docx(self) -> bytes:
        doc = docx.Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("This is the intro paragraph.")
        doc.add_heading("Details", level=2)
        doc.add_paragraph("More detailed content here.")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    async def test_valid_docx(self):
        data = self._make_docx()
        result = await DocxParser().parse(data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert "intro" in result.text.lower()
        assert len(result.sections) >= 1

    async def test_invalid_bytes_raises(self):
        with pytest.raises(ValidationError):
            await DocxParser().parse(b"not a docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
